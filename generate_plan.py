"""수행 계획서 자동 작성 에이전트 — .docx 생성기 (MVP).

기준 템플릿(.docx)의 표·스타일·서식을 보존한 채, 입력 JSON의 가변 값만
주입해 새 수행 계획서를 생성한다. 상세 매핑은 template-spec.md 참조.

사용법:
    python generate_plan.py \
        --template "01.AgentGo PoC(가치 증명) 수행 계획서_v1.2.docx" \
        --input inputs.example.json \
        --output "out/수행계획서_생성본.docx"

설계 메모(왜):
- 기준 템플릿은 플레이스홀더가 없는 '채워진 완성 문서'다. 따라서 Jinja 치환 대신
  기존 텍스트를 찾아 바꾸는 방식을 쓴다. 고객사명은 템플릿에 여러 표기 변형으로
  등장할 수 있어, 입력으로 받은 변형 목록을 긴 것부터 치환해 부분 치환 사고를 막는다.
- 일정 날짜는 'N주차' 라벨을 파싱해 start_date 기준으로 다시 계산한다.
  (단계 수 자체를 바꾸는 것은 MVP 범위 밖 — template-spec.md 참조)
- 자동 생성·미검증 수치에는 검토 유도를 위해 보고만 하고 본문에 마커는 남기지 않는다
  (제출 문서 오염 방지). 검토 필요 항목은 실행 종료 시 요약 출력한다.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import date, datetime, timedelta
from pathlib import Path

try:
    from docx import Document
    from docx.table import _Cell
except ImportError:
    sys.exit("python-docx가 필요합니다. 먼저: pip install -r requirements.txt")


# --- 공통 헬퍼 ---------------------------------------------------------------

def iter_paragraphs(container):
    """문서/셀 내부의 모든 문단을 (중첩 표 포함) 순회한다.
    병합 셀은 같은 문단을 여러 번 노출하지만, 치환은 센티넬 방식으로 멱등하게
    설계해 중복 처리해도 결과가 변하지 않으므로 별도 dedup을 두지 않는다.
    (id() 기반 dedup은 lxml 임시 프록시의 주소 재사용으로 오탐·누락을 일으켜 금지)"""
    for para in container.paragraphs:
        yield para
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def replace_in_paragraph(paragraph, old: str, new: str) -> int:
    """문단 텍스트에서 old를 new로 치환. 런(run) 분할 문제를 피하려고
    문단 전체 텍스트를 합쳐 치환한 뒤 첫 런에 쓰고 나머지를 비운다.
    (치환 구간 내부의 부분 서식은 첫 런 서식으로 통일됨 — MVP 한계, 검증 단계서 확인)
    반환: 치환 횟수."""
    full = "".join(run.text for run in paragraph.runs)
    if old not in full:
        return 0
    count = full.count(old)
    updated = full.replace(old, new)
    if not paragraph.runs:
        return 0
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""
    return count


def replace_everywhere(doc, old: str, new: str) -> int:
    return sum(replace_in_paragraph(p, old, new) for p in iter_paragraphs(doc))


# --- 가변 항목 주입 ----------------------------------------------------------

def replace_client_name(doc, client_name: str, variants: list[str]) -> dict[str, int]:
    """고객사명 변형 표기를 모두 입력값으로 치환.
    client_name 자체가 짧은 변형을 부분 문자열로 포함하면 직접 치환 시
    'creCREinc'처럼 누적 오염이 생긴다. 이를 막기 위해 모든 변형을 먼저 고유
    센티넬로 바꾼 뒤(긴 변형 우선), 센티넬을 client_name으로 한 번에 복원한다.

    반환: {변형: 치환건수}. 호출부가 0건 변형(오타·미사용 표기)과 총 0건(치환
    실패로 이전 고객사명이 문서에 잔존하는 위험)을 구분해 경고/중단하도록 한다."""
    sentinel = "CLIENT_NAME"
    per_variant: dict[str, int] = {}
    for variant in sorted(set(variants), key=len, reverse=True):
        per_variant[variant] = replace_everywhere(doc, variant, sentinel)
    replace_everywhere(doc, sentinel, client_name)
    return per_variant


def set_cover_and_revision(doc, today: date, author: str, approver: str,
                           template_authors: list[str]) -> list[str]:
    """표지 작성일/버전, 개정 이력을 신규 문서 기준으로 정리.
    개정 이력은 v1.0 한 행만 남기고 이전 작성자/승인자(개인정보)를 제거한다."""
    notes = []
    today_str = today.isoformat()

    # 표지 작성일: 'YYYY-MM-DD' 형태를 오늘로
    date_pat = re.compile(r"\d{4}-\d{2}-\d{2}")
    cover_done = False
    for table in doc.tables:
        for row in table.rows:
            label = "".join(c.text for c in row.cells)
            if "작성일" in label:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if date_pat.search("".join(r.text for r in p.runs)):
                            txt = "".join(r.text for r in p.runs)
                            p.runs[0].text = date_pat.sub(today_str, txt)
                            for r in p.runs[1:]:
                                r.text = ""
                            cover_done = True
            if "버  전" in label or re.search(r"버\s*전", label):
                for cell in row.cells:
                    if re.fullmatch(r"\s*\d+\.\d+\s*", cell.text):
                        for p in cell.paragraphs:
                            if p.runs:
                                p.runs[0].text = "1.0"
                                for r in p.runs[1:]:
                                    r.text = ""
    if not cover_done:
        notes.append("표지 작성일 칸을 찾지 못함 — 수동 확인 필요")

    # 개정 이력: 헤더(버전/날짜/내용/작성자/승인자) 다음 데이터 행을 1개만 유지
    rev_table = _find_revision_table(doc)
    if rev_table is None:
        notes.append("개정 이력 표를 찾지 못함 — 수동 확인 필요")
        return notes

    header_idx = _revision_header_index(rev_table)
    data_rows = rev_table.rows[header_idx + 1:]
    first = data_rows[0] if data_rows else None
    if first is not None:
        values = ["1.0", today_str, "초본 작성", author or "[작성자]", approver or "[승인자]"]
        for cell, val in zip(first.cells, values):
            _set_cell_text(cell, val)
    # 나머지 데이터 행 삭제(이전 이력 + 개인정보 제거)
    for row in data_rows[1:]:
        row._element.getparent().remove(row._element)

    return notes


def _find_revision_table(doc):
    for table in doc.tables:
        head = "".join(c.text for c in table.rows[0].cells) if table.rows else ""
        if "문서 개정 이력" in head or _has_revision_header(table):
            return table
    return None


def _has_revision_header(table) -> bool:
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if "버전" in cells and "작성자" in cells and "승인자" in cells:
            return True
    return False


def _revision_header_index(table) -> int:
    for i, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if "버전" in cells and "작성자" in cells:
            return i
    return 0


def _set_cell_text(cell: _Cell, text: str):
    """셀을 단일 문단/텍스트로 정리(첫 문단 서식 유지, 여분 문단 제거).
    날짜가 여러 문단에 흩어진 일정 셀의 옛 값 잔존을 막는다."""
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_solution_paragraph(doc, solution_config: list[str]) -> bool:
    """1장 '검증 대상 솔루션' 자유 문단을 입력값으로 재작성."""
    for p in iter_paragraphs(doc):
        text = p.text
        if "검증 대상 솔루션" in text:
            new = "검증 대상 솔루션(PoC 한정 제공): " + ", ".join(solution_config)
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            return True
    return False


def insert_client_contact(doc, client_pm: str) -> bool:
    """7장 비상연락망(클로잇 측만 존재)에 고객 측 연락처 줄을 신설 추가."""
    for p in iter_paragraphs(doc):
        if "비상 연락망" in p.text:
            new_p = p.insert_paragraph_before("고객사 비상 연락망: " + client_pm)
            new_p.style = p.style
            return True
    return False


# --- 일정 날짜 재계산 --------------------------------------------------------

WEEK_LABEL = re.compile(r"(\d+)\s*주차(?:\s*~\s*(\d+)\s*주차)?")


def _friday_of_week(start: date, week_no: int) -> date:
    """start_date가 속한 주를 1주차로 보고 week_no째 주의 금요일(주 마지막 근무일)."""
    monday_week1 = start - timedelta(days=start.weekday())
    monday = monday_week1 + timedelta(weeks=week_no - 1)
    return monday + timedelta(days=4)  # 금요일


def rebuild_schedule_dates(doc, start: date, total_weeks: int) -> list[str]:
    """일정 표의 'N주차' 라벨을 start_date 기준으로 다시 계산.
    단계 그룹(예: 2주차~3주차)은 유지하고 날짜만 재산출한다."""
    notes = []
    sched = _find_schedule_table(doc)
    if sched is None:
        return ["일정 표를 찾지 못함 — 날짜 수동 확인 필요"]

    for row in sched.rows:
        first_cell = row.cells[0]
        text = first_cell.text
        m = WEEK_LABEL.search(text)
        if not m:
            continue
        start_wk = int(m.group(1))
        end_wk = int(m.group(2)) if m.group(2) else start_wk
        if end_wk > total_weeks:
            notes.append(f"'{text.strip()[:20]}…' 단계가 총 {total_weeks}주차를 초과 — 확인 필요")
        end_date = _friday_of_week(start, end_wk)
        if start_wk == end_wk:
            label = f"{start_wk}주차(~{end_date.month}/{end_date.day})"
        else:
            s_date = _friday_of_week(start, start_wk) - timedelta(days=4)  # 시작주 월요일
            label = f"{start_wk}주차~{end_wk}주차({s_date.month}/{s_date.day}~{end_date.month}/{end_date.day})"
        # 일정(첫) 칸은 날짜 정보만 담으므로 셀 전체를 새 라벨로 정리
        # (옛 날짜가 여러 문단에 흩어져 남는 문제 방지)
        _set_cell_text(first_cell, label)
    return notes


def _find_schedule_table(doc):
    for table in doc.tables:
        head = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        if any("WBS" in h or "수행 항목" in h for h in head) and any("산출물" in h for h in head):
            return table
    return None


# --- 선택 입력(B) 오버라이드 -------------------------------------------------

def apply_optional_overrides(doc, data: dict) -> None:
    os_val = data.get("infra_os")
    if os_val:
        replace_everywhere(doc, "1안_Ubuntu 22.04 LTS, 2안_Rocky Linux 8 또는 9", os_val)
    # KPI/회의주기 등은 표기 변형이 커 단순 치환이 위험 → MVP에서는 보고만.


# --- 메인 --------------------------------------------------------------------

# A그룹(필수) / B그룹(선택). 선택값은 없으면 템플릿 기본값을 쓰므로 차단하지 않고 알린다.
REQUIRED_A = ["client_name", "start_date", "total_weeks", "poc_purpose",
              "verify_tasks", "integration_targets", "solution_config", "client_pm"]
OPTIONAL_B = ["kpi_targets", "infra_os", "use_nas", "meeting_cycle"]

# draft_inputs.py가 모르는 값에 남기는 미확정 표시. 필수값에 이게 있으면 '안 채운 것'으로 본다.
PLACEHOLDER = "[확인 필요]"


def load_inputs(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _is_empty(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        return not v.strip()
    if isinstance(v, (list, dict)):
        return len(v) == 0
    return False


def _has_placeholder(v) -> bool:
    """미확정 표시('[확인 필요]')가 (중첩 포함) 들어 있는지."""
    if isinstance(v, str):
        return PLACEHOLDER in v
    if isinstance(v, dict):
        return any(_has_placeholder(x) for x in v.values())
    if isinstance(v, list):
        return any(_has_placeholder(x) for x in v)
    return False


def validate_inputs(data: dict) -> tuple[list[str], list[str]]:
    """입력값을 점검해 (차단 오류, 비차단 알림)을 돌려준다.
    한 번에 모든 문제를 모아 보고해, 고치고 또 막히는 일을 줄인다."""
    errors: list[str] = []
    warnings: list[str] = []

    # 필수(A): 존재 + 비어있지 않음 + 미확정('[확인 필요]') 아님
    for k in REQUIRED_A:
        v = data.get(k)
        if _is_empty(v):
            errors.append(f"필수값 '{k}' 누락(빈 값)")
        elif _has_placeholder(v):
            errors.append(f"필수값 '{k}' 가 미확정('{PLACEHOLDER}') 상태 — 실제 값으로 채우세요")

    # 형식/타입 점검 (placeholder가 아닌 실제 값에 대해서만)
    sd = data.get("start_date")
    if isinstance(sd, str) and sd.strip() and PLACEHOLDER not in sd:
        try:
            datetime.strptime(sd, "%Y-%m-%d")
        except ValueError:
            errors.append(f"start_date '{sd}' 형식 오류 — YYYY-MM-DD 여야 함")

    tw = data.get("total_weeks")
    if not _is_empty(tw) and not _has_placeholder(tw):
        try:
            if int(tw) < 1:
                errors.append(f"total_weeks 는 1 이상의 정수여야 함 (현재: {tw})")
        except (ValueError, TypeError):
            errors.append(f"total_weeks 가 정수가 아님 (현재: {tw!r})")

    for k in ("verify_tasks", "solution_config"):
        v = data.get(k)
        if v is not None and not isinstance(v, list):
            errors.append(f"{k} 는 목록(list)이어야 함 (현재 타입: {type(v).__name__})")
    it = data.get("integration_targets")
    if it is not None and not isinstance(it, dict):
        errors.append(f"integration_targets 는 객체(object)여야 함 (현재 타입: {type(it).__name__})")

    # 선택(B): 없으면 템플릿 기본값 사용 — 알림(비차단)
    missing_opt = [k for k in OPTIONAL_B if _is_empty(data.get(k))]
    if missing_opt:
        warnings.append("선택값 미입력 → 템플릿/기본값 사용: " + ", ".join(missing_opt))
    for k in OPTIONAL_B:
        if not _is_empty(data.get(k)) and _has_placeholder(data.get(k)):
            warnings.append(f"선택값 '{k}' 가 미확정('{PLACEHOLDER}') — 협의값으로 보정 권장")

    return errors, warnings


# --- 변경 이력 / 재생성 ------------------------------------------------------
# 재생성 시 '직전 입력'과 현재 입력의 차이를 보여주고, 출력 문서별 이력 파일에
# 변경 로그를 누적한다. 비교 대상은 docx가 아니라 입력 JSON이다(의미가 명확하고
# 로컬에만 머물러 안전). 이력 파일은 입력값 스냅샷을 담으므로 출력과 같은 위치
# (out/ → .gitignore)에 두어 커밋되지 않게 한다.

_MISSING = object()


def _is_noise_key(key: str) -> bool:
    """주석·설명용 키(_comment, *_note)는 변경 비교에서 제외한다."""
    return key == "_comment" or key.endswith("_note") or key.endswith("_comment")


def history_path(output: Path) -> Path:
    return output.with_suffix(output.suffix + ".history.json")


def load_history(hist_path: Path) -> dict:
    if hist_path.exists():
        try:
            return json.loads(hist_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}  # 손상된 이력은 무시하고 새로 시작(생성 자체는 막지 않음)
    return {}


def diff_inputs(old: dict, new: dict) -> list[str]:
    """두 입력 스냅샷의 차이를 사람이 읽을 한국어 변경 목록으로 만든다."""
    changes: list[str] = []

    def fmt(v):
        return json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else str(v)

    def walk(prefix, o, n):
        if prefix and _is_noise_key(prefix.split(".")[-1]):
            return
        if isinstance(o, dict) or isinstance(n, dict):
            o = o if isinstance(o, dict) else {}
            n = n if isinstance(n, dict) else {}
            for k in sorted(set(o) | set(n)):
                walk(f"{prefix}.{k}" if prefix else k,
                     o.get(k, _MISSING), n.get(k, _MISSING))
        elif isinstance(o, list) or isinstance(n, list):
            if o != n:
                ol = o if isinstance(o, list) else []
                nl = n if isinstance(n, list) else []
                added = [x for x in nl if x not in ol]
                removed = [x for x in ol if x not in nl]
                parts = []
                if added:
                    parts.append("추가 " + ", ".join(fmt(x) for x in added))
                if removed:
                    parts.append("제거 " + ", ".join(fmt(x) for x in removed))
                changes.append(f"{prefix}: " + ("; ".join(parts) if parts else "순서/내용 변경"))
        elif o != n:
            if o is _MISSING:
                changes.append(f"{prefix} 추가: {fmt(n)}")
            elif n is _MISSING:
                changes.append(f"{prefix} 제거")
            else:
                changes.append(f"{prefix}: {fmt(o)} → {fmt(n)}")

    walk("", old, new)
    return changes


def append_history(hist_path: Path, output: Path, data: dict, changes: list[str]) -> None:
    """이번 재생성을 이력에 기록. 다음 diff를 위해 최신 입력 스냅샷도 보관한다."""
    revisions = load_history(hist_path).get("revisions", [])
    revisions.append({
        "at": datetime.now().isoformat(timespec="seconds"),
        "changes": changes or ["(입력 변경 없음)"],
    })
    hist_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {"document": str(output), "last_inputs": data, "revisions": revisions}
    hist_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
                         encoding="utf-8")


def print_history(hist_path: Path) -> None:
    revisions = load_history(hist_path).get("revisions", [])
    if not revisions:
        print(f"이력 없음: {hist_path}")
        return
    print(f"재생성 이력 — 총 {len(revisions)}회")
    for i, rev in enumerate(revisions, 1):
        print(f"\n[{i}] {rev.get('at', '')}")
        for c in rev.get("changes", []):
            print(f"    - {c}")


def main():
    # Windows 콘솔(cp949)에서 한글·em-dash 출력이 깨지지 않도록 UTF-8로 고정.
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass

    ap = argparse.ArgumentParser(description="수행 계획서 .docx 자동 생성기 (MVP)")
    ap.add_argument("--template")
    ap.add_argument("--input")
    ap.add_argument("--output", required=True)
    ap.add_argument("--allow-no-client-name", action="store_true",
                    help="고객사명 치환 0건이어도 중단하지 않고 경고만 표시 "
                         "(가명 예시·테스트용. 실제 문서 생성에는 쓰지 말 것)")
    ap.add_argument("--show-history", action="store_true",
                    help="이 출력 문서(--output)의 재생성 변경 이력을 출력하고 종료")
    ap.add_argument("--no-history", action="store_true",
                    help="이번 실행을 변경 이력에 기록하지 않음")
    ap.add_argument("--pdf", action="store_true",
                    help="생성된 .docx 를 PDF로도 변환 (MS Word 필요)")
    ap.add_argument("--hwp", action="store_true",
                    help="생성된 .docx 를 HWP로도 변환 (한컴오피스 필요. 보안 승인 창 허용 필요)")
    args = ap.parse_args()

    out_path = Path(args.output)
    hist_path = history_path(out_path)

    # 이력 조회 모드: 생성 없이 과거 변경 이력만 보여주고 종료.
    if args.show_history:
        print_history(hist_path)
        return

    if not args.template or not args.input:
        ap.error("--template 와 --input 은 필수입니다 (--show-history 제외).")

    template_path = Path(args.template)
    if not template_path.exists():
        sys.exit(f"템플릿 없음: {template_path}")

    data = load_inputs(Path(args.input))

    # 빠진 값/형식 점검: 선택값 미입력은 알리고, 필수값 누락·미확정·형식오류는 저장 전에 막는다.
    errors, warnings = validate_inputs(data)
    for w in warnings:
        print(f"- 빠진 값 점검(알림): {w}")
    if errors:
        print("\n❌ 입력 검증 실패 — 다음을 수정한 뒤 다시 실행하세요:", file=sys.stderr)
        for e in errors:
            print(f"  - {e}", file=sys.stderr)
        sys.exit(1)

    start = datetime.strptime(data["start_date"], "%Y-%m-%d").date()
    today = date.today()

    # 직전 생성과의 입력 차이(변경 이력)를 계산해 보여준다.
    prev_inputs = load_history(hist_path).get("last_inputs")
    if prev_inputs is None:
        change_log = ["초기 생성"]
        print("- 변경 이력: 첫 생성(이전 이력 없음)")
    else:
        change_log = diff_inputs(prev_inputs, data)
        if change_log:
            print(f"- 변경 이력: 직전 대비 {len(change_log)}건 변경")
            for c in change_log:
                print(f"    · {c}")
        else:
            print("- 변경 이력: 직전과 입력 동일(내용 변경 없음)")

    doc = Document(str(template_path))
    review_notes: list[str] = []

    # 템플릿에 박힌 고객사명 표기 변형 목록. 입력에서 받고, 없으면 client_name 단일값.
    variants = data.get("_client_name_variants") or [data["client_name"]]
    per_variant = replace_client_name(doc, data["client_name"], variants)
    n = sum(per_variant.values())
    print(f"- 고객사명 치환: {n}건 (" +
          ", ".join(f"{v}:{c}" for v, c in per_variant.items()) + ")")

    # 0건 변형은 표기 오타이거나 템플릿에 없는 표기 — 검토 항목으로만 보고.
    for v, c in per_variant.items():
        if c == 0:
            review_notes.append(
                f"고객사명 변형 '{v}' 이(가) 템플릿에서 한 번도 발견되지 않음"
                " — 표기 오타이거나 불필요한 항목인지 확인")

    # 총 0건은 위험 신호: 치환이 전혀 안 됐다면 템플릿의 '이전 고객사명'이 그대로
    # 남아 개인정보가 유출될 수 있다. 기본은 저장 전에 중단하고, 가명 예시·테스트에
    # 한해 --allow-no-client-name 으로 경고만 남기고 진행한다.
    if n == 0:
        guidance = (
            "고객사명 치환 0건 — 입력한 변형이 템플릿의 실제 표기와 하나도 일치하지 않습니다.\n"
            "  이대로 저장하면 템플릿에 박힌 '이전 고객사명'이 새 문서에 그대로 남아\n"
            "  개인정보·고객사 정보가 유출될 수 있습니다.\n"
            "  → inputs의 _client_name_variants 에 템플릿에 실제로 박힌 표기(모든 변형)를 넣으세요.")
        if args.allow_no_client_name:
            print(f"\n⚠️  경고: {guidance}")
            review_notes.append(
                "고객사명 치환 0건 상태로 강제 진행됨(--allow-no-client-name)"
                " — 출력 문서에 남은 이전 고객사명을 반드시 직접 확인할 것")
        else:
            sys.exit(f"\n❌ 중단(저장 안 함): {guidance}\n"
                     "  (가명 예시로 의도한 동작이면 --allow-no-client-name 옵션으로 진행)")

    review_notes += set_cover_and_revision(
        doc, today, data.get("author", ""), data.get("approver", ""),
        data.get("_template_revision_authors", []))

    if replace_solution_paragraph(doc, data["solution_config"]):
        print("- 검증 대상 솔루션 문단 재작성 완료")
    else:
        review_notes.append("검증 대상 솔루션 문단을 찾지 못함")

    if insert_client_contact(doc, data["client_pm"]):
        print("- 고객 측 비상연락망 신설 완료")
    else:
        review_notes.append("비상연락망 위치를 찾지 못해 고객 연락처 미삽입")

    review_notes += rebuild_schedule_dates(doc, start, int(data["total_weeks"]))
    print("- 일정 날짜 재계산 완료")

    apply_optional_overrides(doc, data)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"\n생성 완료: {out_path}")

    # 변경 이력 기록(저장 성공 후). 다음 재생성의 diff 기준이 될 입력 스냅샷도 함께 보관.
    if not args.no_history:
        append_history(hist_path, out_path, data, change_log)
        print(f"이력 기록: {hist_path}  (--show-history 로 조회)")

    # 출력 포맷 확장: 서식 보존을 위해 새로 렌더링하지 않고 설치된 오피스로 .docx 를 변환.
    # 변환 실패는 검토 항목으로만 남기고, 이미 저장된 .docx 생성은 성공으로 둔다.
    if args.pdf or args.hwp:
        import convert
        if args.pdf:
            try:
                print(f"PDF 변환 완료: {convert.to_pdf(out_path)}")
            except Exception as e:
                review_notes.append(f"PDF 변환 실패({type(e).__name__}: {e}) — .docx 는 정상 생성됨")
        if args.hwp:
            print("  (HWP: 한컴 보안 승인 창이 뜨면 '허용'하세요. 보안 모듈 미등록 시 멈출 수 있음)")
            try:
                print(f"HWP 변환 완료: {convert.to_hwp(out_path)}")
            except Exception as e:
                review_notes.append(f"HWP 변환 실패({type(e).__name__}: {e}) — .docx 는 정상 생성됨")

    # 검토 필요 항목 요약(본문 오염 없이 콘솔로만)
    print("\n[확인 필요 — 제출 전 검토]")
    default_review = [
        "일정 표 WBS 세부 항목/산출물은 템플릿 기준 그대로 — 과제에 맞는지 확인",
        "성공 기준(KPI) 목표 수준은 템플릿 기본값 — 협의 결과 반영 여부 확인",
        "솔루션 구성·연동 대상이 실제 계약 범위와 일치하는지 확인",
    ]
    for note in default_review + review_notes:
        print(f"  - {note}")


if __name__ == "__main__":
    main()
